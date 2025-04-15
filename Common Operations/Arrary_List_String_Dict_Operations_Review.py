from docx import Document
from docx.shared import Inches

def create_common_operations_doc():
    """Creates a Word document describing common operations for arrays, lists, strings, and dictionaries."""

    document = Document()

    document.add_heading('Common Operations in Python Data Structures', level=1)

    # Arrays (using lists as the primary array-like structure)
    document.add_heading('Arrays (Lists)', level=2)
    document.add_paragraph("In Python, lists are commonly used to represent arrays.  While Python has an `array` module, lists are more versatile.")

    operations = [
        ("Accessing an element", "`my_list[index]`", " (O(1)) - Accesses the element at the specified index."),  # Added comma
        ("Appending an element", "`my_list.append(item)`", " (O(1) amortized) - Adds an element to the end of the list."), # Added comma
        ("Inserting an element", "`my_list.insert(index, item)`", " (O(n)) - Inserts an element at a specific index."), # Added comma
        ("Removing an element (by value)", "`my_list.remove(item)`", " (O(n)) - Removes the first occurrence of an element."), # Added comma
        ("Removing an element (by index)", "`my_list.pop(index)`", " (O(n) in general, O(1) for the last element) - Removes and returns the element at the specified index."), # Added comma
        ("Finding the length", "`len(my_list)`", " (O(1)) - Returns the number of elements in the list."), # Added comma
        ("Slicing", "`my_list[start:end:step]`", " (O(k), where k is the size of the slice) - Creates a new list containing a portion of the original list."), # Added comma
        ("Iterating", "`for item in my_list:`", " (O(n)) - Iterates through the elements of the list."), # Added comma
        ("Checking for membership", "`item in my_list`", " (O(n)) - Checks if an element exists in the list."), # Added comma
        ("Sorting", "`my_list.sort()`", " (in-place, O(n log n)) or `sorted(my_list)` (returns a new sorted list, O(n log n)) - Sorts the list."), # Added comma
        ("Reversing", "`my_list.reverse()`", " (in-place, O(n)) or `reversed(my_list)` (returns an iterator, O(n) to iterate) - Reverses the order of elements."), # Added comma
        ("Concatenation", "`list1 + list2`", " (O(n+m) where n and m are the lengths of lists ) - create a new list containing all elements.") # Added comma
    ]

    table = document.add_table(rows=1, cols=3)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Operation'
    hdr_cells[1].text = 'Example'
    hdr_cells[2].text = 'Description/Time Complexity'
    for op, example, desc in operations:
        row_cells = table.add_row().cells
        row_cells[0].text = op
        row_cells[1].text = example
        row_cells[2].text = desc

    # Lists (covered above, so just a brief mention)
    # document.add_heading('Lists', level=2)
    # document.add_paragraph("Lists are covered in the Arrays section above, as they are the primary way to represent arrays in Python.")

    # Strings
    document.add_heading('Strings', level=2)
    document.add_paragraph("Strings are sequences of characters and are immutable.")

    operations = [
        ("Accessing a character", "`my_string[index]`", " (O(1)) - Accesses the character at the specified index."), # Added comma
        ("Slicing", "`my_string[start:end:step]`", " (O(k), where k is the size of the slice) - Creates a new string containing a portion of the original string."), # Added comma
        ("Concatenation", "`string1 + string2`", " (O(n+m), where n and m are string lengths) - Creates a new string by joining two strings."), # Added comma
        ("Finding the length", "`len(my_string)`", " (O(1)) - Returns the number of characters in the string."), # Added comma
        ("Iterating", "`for char in my_string:`", " (O(n)) - Iterates through the characters of the string."), # Added comma
        ("Checking for substring", "`substring in my_string`", " (O(nm) in worst case, but often faster with optimized algorithms like KMP) - Checks if a substring exists in the string."), # Added comma
        ("Converting case", "`my_string.lower()`", " (O(n)), `my_string.upper()` (O(n)) - Converts the string to lowercase or uppercase."), # Added comma
        ("Replacing substrings", "`my_string.replace(old, new)`", " (O(n)) - Creates a new string with replaced substrings."), # Added comma
        ("Splitting", "`my_string.split(separator)`", " (O(n)) - Creates a list of strings by splitting the string at the separator."), # Added comma
        ("Joining", "`separator.join(list_of_strings)`", " (O(n)) - Creates a string by joining elements of a list with a separator."), # Added comma
        ("Finding index of substring", "`string.find(substring)`", "(O(nm)) - Returns the index of the first occurrence of the substring or -1 if it is not found."), # Added comma
         ("Stripping whitespace", "`my_string.strip()`", "  (O(n))- returns a new string with leading and trailing whitespaces removed.") # Added comma
    ]

    table = document.add_table(rows=1, cols=3)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Operation'
    hdr_cells[1].text = 'Example'
    hdr_cells[2].text = 'Description/Time Complexity'
    for op, example, desc in operations:
        row_cells = table.add_row().cells
        row_cells[0].text = op
        row_cells[1].text = example
        row_cells[2].text = desc

    # Dictionaries
    document.add_heading('Dictionaries', level=2)
    document.add_paragraph("Dictionaries store key-value pairs. Keys must be immutable (e.g., strings, numbers, tuples).")

    operations = [
        ("Accessing a value", "`my_dict[key]`", " (O(1) average, O(n) worst case) - Accesses the value associated with the key."), # Added comma
        ("Adding or updating a key-value pair", "`my_dict[key] = value`", " (O(1) average, O(n) worst case) - Adds a new key-value pair or updates the value for an existing key."), # Added comma
        ("Removing a key-value pair", "`del my_dict[key]`", " (O(1) average, O(n) worst case) - Removes the key-value pair."), # Added comma
        ("Checking for key existence", "`key in my_dict`", " (O(1) average, O(n) worst case) - Checks if a key exists in the dictionary."), # Added comma
        ("Getting all keys", "`my_dict.keys()`", " (O(1) - returns a view object) - Returns a view object of all keys."), # Added comma
        ("Getting all values", "`my_dict.values()`", " (O(1) - returns a view object) - Returns a view object of all values."), # Added comma
        ("Getting all key-value pairs", "`my_dict.items()`", " (O(1) - returns a view object) - Returns a view object of all key-value pairs as tuples."), # Added comma
        ("Iterating", "`for key, value in my_dict.items():`", " (O(n)) - Iterates through the key-value pairs."), # Added comma
        ("Finding the length (number of key-value pairs)", "`len(my_dict)`", " (O(1)) - Returns the number of key-value pairs."), # Added comma
        ("Removing and return a value", "`my_dict.pop(key)`", " (O(1) average, O(n) worst-case) - Returns value associated with the key. Key value pair also removed. Raises KeyError if Key not found."), # Added comma
        ("Get a value, or default", "`my_dict.get(key, default_value)`", " (O(1) average, O(n) worst case) - returns the value associated with key, or default_value if the key isn't present.") # Added comma
    ]

    table = document.add_table(rows=1, cols=3)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Operation'
    hdr_cells[1].text = 'Example'
    hdr_cells[2].text = 'Description/Time Complexity'
    for op, example, desc in operations:
        row_cells = table.add_row().cells
        row_cells[0].text = op
        row_cells[1].text = example
        row_cells[2].text = desc

    document.save('common_data_structure_operations.docx')

create_common_operations_doc()
print("Word document 'common_data_structure_operations.docx' created successfully.")