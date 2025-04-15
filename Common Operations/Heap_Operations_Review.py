from docx import Document
from docx.shared import Inches

def create_heap_operations_doc():
    """Creates a Word document explaining heapq operations."""

    document = Document()

    document.add_heading('Python heapq Module: Heap Operations', level=1)

    document.add_paragraph(
        "Python's `heapq` module provides an implementation of the heap queue algorithm (min-heap). "
        "Here's a list of the widely used heap operations provided by `heapq`, along with brief explanations:"
    )

    # heappush
    document.add_heading('heapq.heappush(heap, item)', level=2)
    document.add_paragraph("Purpose: Pushes the `item` onto the `heap`, maintaining the heap invariant (parent node is less than or equal to its children in a min-heap).")
    document.add_paragraph("Explanation: Inserts a new element into the heap and then \"heapifies\" upwards (sifts up) to ensure the heap property is preserved.")

    # heappop
    document.add_heading('heapq.heappop(heap)', level=2)
    document.add_paragraph("Purpose: Pops and returns the smallest item from the `heap`, maintaining the heap invariant.")
    document.add_paragraph("Explanation: Removes the root element (smallest item) and then replaces it with the last element of the heap. It then \"heapifies\" downwards (sifts down) to restore the heap property.")

    # heappushpop
    document.add_heading('heapq.heappushpop(heap, item)', level=2)
    document.add_paragraph("Purpose: Pushes `item` on the heap and then pops and returns the smallest item. More efficient than a separate `heappush()` followed by a `heappop()`.")
    document.add_paragraph("Explanation:  Combines the push and pop operations.  It's faster because it only needs to re-heapify once.")

    # heapreplace
    document.add_heading('heapq.heapreplace(heap, item)', level=2)
    document.add_paragraph("Purpose: Pops and returns the smallest item from the `heap`, and then pushes the new `item`.  The heap size doesn't change. More efficient than `heappop()` followed by `heappush()`, especially when the `item` might be smaller than the popped element.")
    document.add_paragraph("Explanation: This is essentially a `heappop()` followed by a `heappush()`, but avoids an unnecessary \"sift up\" if the new `item` is larger than all existing heap elements.")

    # heapify
    document.add_heading('heapq.heapify(x)', level=2)
    document.add_paragraph("Purpose: Transforms list `x` into a heap *in-place*, in linear time (O(n)).")
    document.add_paragraph("Explanation: Rearranges the elements of the list `x` to satisfy the heap property.  This is often used to initialize a heap from an existing list.  Crucially, this operates *in-place*, modifying the original list.")

    # nlargest
    document.add_heading('heapq.nlargest(n, iterable, key=None)', level=2)
    document.add_paragraph("Purpose: Returns a list with the `n` largest elements from the `iterable`. `key` specifies a function of one argument that is used to extract a comparison key from each element in `iterable` (similar to `sorted()`).")
    document.add_paragraph("Explanation: Uses a heap to efficiently find the largest elements.  More efficient than sorting the entire iterable when `n` is significantly smaller than the size of `iterable`.")
    # nsmallest
    document.add_heading('heapq.nsmallest(n, iterable, key=None)', level=2)
    document.add_paragraph("Purpose: Returns a list with the `n` smallest elements from the `iterable`. `key` has the same function as in `nlargest`.")
    document.add_paragraph("Explanation:  Similar to `nlargest`, but finds the smallest elements.  Also more efficient than sorting the entire iterable when `n` is small.")

    # Important Considerations
    document.add_heading('Important Considerations (and things *not* directly in `heapq`)', level=1)

    document.add_paragraph(
        "**Decreasing a Key / Changing Priority:** The `heapq` module does *not* provide a direct way to decrease the key (priority) of an existing element in the heap, or to change the priority of element. "
        "Implementing this efficiently requires additional data structures (often a dictionary mapping elements to their heap indices). "
        "This is a common requirement in algorithms like Dijkstra's shortest path. "
        "A common workaround is to push a new tuple `(new_priority, item)` onto the heap, even if `item` already exists. "
        "You then need logic to ignore \"stale\" entries when they are popped (checking if the popped item's priority is the *current* lowest known priority)."
    )

    document.add_paragraph(
        "**Accessing Elements by Index (Without Popping):** `heapq` works on Python lists. "
        "While you *can* access elements by index (like `heap[3]`), doing so *does not* guarantee any particular order except for the smallest element (at `heap[0]`). "
        "Directly accessing elements by index is generally *not* part of typical heap usage and can break the heap property if you modify the list directly.  "
        "The heap is designed for efficient retrieval of the *smallest* element, not arbitrary element access."
    )

    document.add_paragraph(
        "**Max-Heaps:**  `heapq` implements a *min-heap*.  "
        "To simulate a max-heap (where you want to retrieve the *largest* element), a common technique is to negate the values you insert and then negate them again when you retrieve them.  "
        "Alternatively, you can use a custom comparison function (using the `key` argument in functions like `nlargest` and `nsmallest` which can use a custom compare function)."
    )

    document.save('heap_operations.docx')

create_heap_operations_doc()

print("Word document 'heap_operations.docx' created successfully.")